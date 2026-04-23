"""
Generate the Shibumi Nadlan brief document (Word .docx, Hebrew / RTL).

This is a one-shot script; it writes the brief to the project folder.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from pathlib import Path


# ---------- helpers ----------

def set_rtl(paragraph):
    """Force a paragraph to render right-to-left (Hebrew)."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def set_run_rtl(run):
    """Mark run as RTL-complex-script."""
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def add_p(doc, text="", *, style=None, size=11, bold=False, italic=False,
          color=None, align=WD_ALIGN_PARAGRAPH.RIGHT, rtl=True, font="Assistant",
          space_before=None, space_after=None, after_pt=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.alignment = align
    if rtl:
        set_rtl(p)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    elif after_pt is not None:
        p.paragraph_format.space_after = Pt(after_pt)

    if text:
        run = p.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = RGBColor(*color)
        # Set East Asian / complex script font too so Word applies it to Hebrew
        rFonts = run._element.get_or_add_rPr().find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            run._element.get_or_add_rPr().append(rFonts)
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)
        rFonts.set(qn("w:cs"),    font)
        if rtl:
            set_run_rtl(run)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 22, 2: 15, 3: 12}
    return add_p(
        doc, text,
        size=sizes.get(level, 12),
        bold=True,
        color=(20, 23, 28),
        space_before=14, space_after=6,
        font="Assistant",
    )


def add_eyebrow(doc, text):
    return add_p(
        doc, text.upper(),
        size=8, bold=True,
        color=(154, 133, 81),  # brass
        space_before=18, space_after=2,
        font="Assistant",
    )


def add_field(doc, label, placeholder="_________________________________________________"):
    add_p(doc, label, size=11, bold=True, color=(36, 40, 48), space_before=8, space_after=2)
    add_p(doc, placeholder, size=11, color=(120, 120, 130), space_after=4)


def add_checkbox_line(doc, text):
    add_p(doc, f"☐  {text}", size=11, color=(36, 40, 48), space_after=2)


def add_rule(doc):
    """Horizontal rule via a paragraph border."""
    p = add_p(doc, "", size=1, space_before=8, space_after=8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9B483")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_default_rtl(doc):
    """Make the whole document render RTL by default."""
    settings = doc.settings.element
    # Add <w:bidi/> to default paragraph properties if possible
    # and set sectPr bidi.
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = OxmlElement("w:bidi")
        sectPr.append(bidi)


# ---------- content ----------

def build():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    set_default_rtl(doc)

    # ---- Cover ----
    add_p(doc, "SHIBUMI NADLAN", size=9, bold=True, color=(154, 133, 81),
          font="Assistant", align=WD_ALIGN_PARAGRAPH.RIGHT, after_pt=4)

    add_p(doc, "בריף לבעל החברה — חומרים ומידע נדרשים לאתר החדש",
          size=26, bold=True, color=(20, 23, 28),
          font="Assistant", after_pt=6)

    add_p(doc, "מסמך פנימי · לא להעברה לגורם שלישי ללא אישור",
          size=9, italic=True, color=(138, 144, 156), after_pt=2)

    add_rule(doc)

    add_p(doc, "מייקל שלום,", size=12, after_pt=4)
    add_p(
        doc,
        "להשלמת האתר החדש של שיבומי נדל״ן אנחנו צריכים מספר חומרים וקטעי מידע. "
        "המסמך הזה מרכז את כל מה שחסר — חלקו חובה, חלקו רצוי מאוד, וחלקו אופציונלי "
        "אבל יעשה הבדל דרמטי בתוצאה הסופית.",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_p(
        doc,
        "מה שלא יהיה בידינו — אנחנו נכתוב בעצמנו / נעצב בעצמנו, ואז תעבור אישור לפני שעולה לאוויר. "
        "אין לחץ לענות על הכל בבת אחת; עדיף לשלוח מדור מדור.",
        size=11, color=(60, 63, 75), after_pt=10,
    )

    add_p(doc, "תודה,", size=11, color=(36, 40, 48), after_pt=1)
    add_p(doc, "צוות פיתוח שיבומי", size=11, bold=True, color=(20, 23, 28), after_pt=16)

    # =================================================================
    # SECTION 1 — BRAND IDENTITY
    # =================================================================
    add_eyebrow(doc, "חלק 1 · זהות המותג")
    add_heading(doc, "לוגו, שם וטאגליין", level=1)
    add_rule(doc)

    add_heading(doc, "1.1  לוגו", level=2)
    add_p(
        doc,
        "אם יש לוגו רשמי של החברה, נבקש אותו בפורמט הכי איכותי שיש בידך "
        "(עדיף SVG או AI; אחרת PNG ברקע שקוף, מעל 2000px).",
        size=11, color=(60, 63, 75), after_pt=6,
    )
    add_p(doc, "סמן את המצב הרלוונטי:", size=11, bold=True, after_pt=4)
    add_checkbox_line(doc, "יש לי קובץ לוגו מקצועי ואשלח (SVG / AI / PNG גבוה)")
    add_checkbox_line(doc, "יש לי לוגו, אבל בקובץ ישן / באיכות נמוכה — אשלח מה שיש ותעשו upgrade")
    add_checkbox_line(doc, "אין לי לוגו — אתם רשאים לעצב מונוגרם חדש «渋 · Shibumi» לאישור שלי")
    add_p(doc, "הערות:", size=10, color=(120,120,130), space_before=4, after_pt=1)
    add_field(doc, "")

    add_heading(doc, "1.2  שם הדומיין הרשמי", level=2)
    add_field(doc, "שם העסק הרשום (לטובת footer / יצוגיות):")
    add_field(doc, "אם יש ח.פ / ע.מ — נציג רק בדף אודות מקוצר:")

    add_heading(doc, "1.3  טאגליין (משפט מזהה קצר)", level=2)
    add_p(
        doc,
        "המשפט שמופיע מתחת לשם החברה בכותרת ובפוטר. צריך להיות קצר, אישי ומזהה.",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_p(doc, "שלוש אפשרויות שהכנו — סמן את המועדף או הצע חלופה:", size=11, bold=True, after_pt=4)
    add_checkbox_line(doc, "חברת בוטיק לנדל״ן יוקרה · מאז 1997")
    add_checkbox_line(doc, "Life Style Real Estate — הבית שמתאים לך, לא רק למפרט")
    add_checkbox_line(doc, "נדל״ן שקט, אישי, ייחודי")
    add_field(doc, "הצעה משלך:")

    add_heading(doc, "1.4  סיפור השם «Shibumi / 渋味»", level=2)
    add_p(
        doc,
        "«Shibumi» הוא מושג יפני שמתאר יופי שקט ומאופק — בדיוק המהות של החברה. "
        "לטובת דף אודות ודף הבית, אנחנו רוצים גרסה שלך (3–5 שורות) לסיפור הבחירה בשם.",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_p(doc, "אם אין לך נוסח מוכן — סמן ונכתוב טיוטה שתאשר:", size=11, bold=True, after_pt=4)
    add_checkbox_line(doc, "יש לי נוסח — אשלח בנפרד")
    add_checkbox_line(doc, "כתבו אתם טיוטה (2–3 גרסאות) לאישור שלי")
    add_field(doc, "הסיפור שלך (אם יש):")

    # =================================================================
    # SECTION 2 — THE COMPANY
    # =================================================================
    add_eyebrow(doc, "חלק 2 · מידע על החברה")
    add_heading(doc, "פרטים שמופיעים בדף אודות", level=1)
    add_rule(doc)

    add_heading(doc, "2.1  פרטים של מייקל ישראלסטאם", level=2)
    add_field(doc, "שם מלא כפי שצריך להופיע באתר:")
    add_field(doc, "תפקיד רשמי (מייסד ומנכ״ל / מייסד / Owner & Founder):")
    add_field(doc, "שנת לידה / רקע קצר (משפט או שניים — אופציונלי):")

    add_heading(doc, "2.2  סיפור הקמת החברה (1997)", level=2)
    add_p(
        doc,
        "2–3 פסקאות קצרות על הסיפור שמאחורי החברה. מה הוביל להקמה, מה השונות, "
        "איך שמתם את Life Style במרכז. אם מעדיף, נוציא מראיון אישי של 15 דק׳ טלפונית.",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_checkbox_line(doc, "אכתוב / אשלח סיפור מוכן")
    add_checkbox_line(doc, "מעדיף ראיון טלפוני של 15 דקות — הפעילו אותי ותכתבו אתם לאישור")
    add_field(doc, "הערות:")

    add_heading(doc, "2.3  מספרים וסטטיסטיקות (מרשים ליוקרה)", level=2)
    add_p(
        doc,
        "מספרים גדולים על דף הבית יוצרים רושם של ותק ומקצועיות. "
        "סמן מה שמוכן להופיע (גם אם בערך):",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_field(doc, "שנות פעילות (1997 → היום = 28):")
    add_field(doc, "מספר משפחות / לקוחות שליוות (בערך):")
    add_field(doc, "היקף עסקאות מצטבר (בערך, אם מוכן לשתף):")
    add_field(doc, "שיא גודל עסקה (בש״ח, בערך):")
    add_field(doc, "נכסים פעילים כיום במלאי:")
    add_field(doc, "גדול ביותר בדונם שנמכר (למשל: נחלה של 40 דונם):")

    add_heading(doc, "2.4  תחומי התמחות — דיוק", level=2)
    add_p(
        doc,
        "מהאתר הקיים זיהינו את הסוגים הבאים. סמן ששבע הקטגוריות הבאות מדויקות, "
        "והוסף / הורד לפי הצורך:",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    for item in [
        "אחוזות מפוארות",
        "נחלות ומשקים חקלאיים (עד 30 דונם)",
        "חוות פרטיות",
        "משקי עזר",
        "וילות ובתי יוקרה",
        "דירות יוקרה",
        "פנטהאוזים",
        "נכסים קו ראשון לים",
    ]:
        add_checkbox_line(doc, item)
    add_field(doc, "להוסיף:")
    add_field(doc, "להסיר:")

    add_heading(doc, "2.5  אזורי פעילות עיקריים", level=2)
    for item in [
        "שרון וחוף בית ינאי (בית ינאי, מכמורת, חופית, כפר ויתקין, ביתן אהרון)",
        "הרצליה פיתוח ותל אביב",
        "עמק יזרעאל וצפון",
        "ארסוף",
        "נווה ירק / בני ציון / פסטורלי בשרון",
        "עמיקם",
        "רשפון",
        "יפו",
    ]:
        add_checkbox_line(doc, item)
    add_field(doc, "אזורים נוספים שחשוב להציג:")
    add_field(doc, "אזורים שאנחנו לא פועלים בהם (חשוב לא להציג):")

    # =================================================================
    # SECTION 3 — CONTACT
    # =================================================================
    add_eyebrow(doc, "חלק 3 · יצירת קשר")
    add_heading(doc, "פרטי התקשרות לאתר", level=1)
    add_rule(doc)

    add_p(doc, "המספרים שזיהינו מהאתר הקיים — נא לאשר / לתקן:",
          size=11, color=(60, 63, 75), after_pt=6)

    add_p(doc, "טלפון ראשי:  050-733-0090", size=11, bold=True, after_pt=0)
    add_checkbox_line(doc, "נכון ופעיל")
    add_checkbox_line(doc, "יש לעדכן — המספר הנכון:  ___________________")

    add_p(doc, "WhatsApp:  +972-50-733-0090", size=11, bold=True, space_before=6, after_pt=0)
    add_checkbox_line(doc, "נכון ופעיל (זה גם המספר שמופיע כ־WhatsApp בכל דף)")
    add_checkbox_line(doc, "יש לעדכן — המספר הנכון:  ___________________")

    add_heading(doc, "3.1  ערוצי קשר נוספים", level=2)
    add_field(doc, "כתובת דוא״ל רשמית:")
    add_field(doc, "כתובת משרד (אם יש, אחרת נשאיר ללא):")
    add_field(doc, "שעות פעילות (למשל: א'–ה' 09:00–18:00):")
    add_field(doc, "קישור Instagram:")
    add_field(doc, "קישור Facebook:")
    add_field(doc, "קישור LinkedIn של מייקל:")

    add_heading(doc, "3.2  ביקורות גוגל", level=2)
    add_p(
        doc,
        "הקישור הבא מהאתר הקיים משמש לכפתור «ביקורות בגוגל» — נא לאשר שהוא תקין:",
        size=11, color=(60, 63, 75), after_pt=2,
    )
    add_p(
        doc,
        "https://www.google.com/search?q=שיבומי+נדלן#lrd=0x151d14cb8104a827:0x655f1df93c88cd7d",
        size=9, font="Consolas", color=(90, 97, 110), after_pt=4, rtl=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    add_checkbox_line(doc, "תקין ופעיל")
    add_checkbox_line(doc, "יש להשתמש בקישור אחר:")
    add_field(doc, "")

    # =================================================================
    # SECTION 4 — VISUAL ASSETS
    # =================================================================
    add_eyebrow(doc, "חלק 4 · חומרים ויזואליים")
    add_heading(doc, "תמונות, וידאו ומסמכים ויזואליים", level=1)
    add_rule(doc)

    add_heading(doc, "4.1  תמונת פורטרט של מייקל (רצוי מאוד)", level=2)
    add_p(
        doc,
        "תמונה של מייקל משדרגת משמעותית את דף האודות ויוצרת אמון. "
        "עדיף צילום שחור-לבן, פרופיל או חצי פרופיל, רקע נקי. "
        "הצילום יכול להיות גם בבית / בנכס — לא חייב סטודיו.",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_checkbox_line(doc, "יש לי תמונה מקצועית — אשלח")
    add_checkbox_line(doc, "יש לי רק תמונה יומיומית — אשלח, אם מתאים נשתמש בה")
    add_checkbox_line(doc, "מעדיף בלי תמונה אישית — לעבוד עם טיפוגרפיה וציטוט")
    add_field(doc, "הערות:")

    add_heading(doc, "4.2  לוגו / מונוגרם באיכות SVG", level=2)
    add_p(
        doc,
        "חזרה על נקודה חשובה. ללא קובץ וקטורי, הלוגו נראה מטושטש בריזולוציות גבוהות. "
        "אם אין — אנחנו נעצב מונוגרם בהתבסס על סימן הקאנג׳י 渋 (Shibumi) ונגיש לאישור.",
        size=11, color=(60, 63, 75), after_pt=4,
    )

    add_heading(doc, "4.3  3 נכסים פלאג׳שיפ ל־Hero", level=2)
    add_p(
        doc,
        "בדף הבית של האתר יש «hero» גדול שמציג שלוש תמונות מרשימות (אחת מהן נחתכת בתוך אותיות השם). "
        "בחר שלושה נכסים מהפורטפוליו שמייצגים הכי טוב את החברה (פעילים / נמכרים — זה לא משנה לטובת הצגה ויזואלית):",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_field(doc, "נכס #1 (שם / מיקום):")
    add_field(doc, "נכס #2 (שם / מיקום):")
    add_field(doc, "נכס #3 (שם / מיקום):")
    add_p(
        doc,
        "(אם אין העדפה — אנחנו נבחר את שלושת הנכסים עם הכי הרבה תמונות איכותיות ונעביר לאישור.)",
        size=10, italic=True, color=(120, 120, 130), after_pt=2,
    )

    add_heading(doc, "4.4  תמונות חסרות / איכות נמוכה", level=2)
    add_p(
        doc,
        "במלאי הנכסים יש חלקם עם רק תמונה אחת. אם יש בידך תמונות נוספות לנכסים הבאים, נשלב אותן. "
        "רשימת נכסים עם מעט תמונות נעביר אליך בקובץ נפרד — עד לקבלתן נעבוד עם מה שיש.",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_checkbox_line(doc, "אעדכן / אשלח תמונות חסרות לפי הרשימה שתעבירו")
    add_checkbox_line(doc, "העבירו את הרשימה ואבדוק מה שיש")

    add_heading(doc, "4.5  וידאו / Drone footage (חזק מאוד ב־Hero)", level=2)
    add_p(
        doc,
        "וידאו קצר של 8–15 שניות (ללא קול) משדרג את תחושת היוקרה בצורה דרמטית. "
        "אם יש Drone footage של נחלה גדולה / בית על הים / נכס מיוחד — נשבץ אותו ב־Hero במקום תמונה.",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_checkbox_line(doc, "יש לי וידאו מתאים — אשלח")
    add_checkbox_line(doc, "אין לי כרגע — נסתפק בתמונה סטטית")
    add_field(doc, "הערות:")

    add_heading(doc, "4.6  כתבות / מדיה שכיסו את שיבומי", level=2)
    add_p(
        doc,
        "אם יש כתבות בעיתונות, כתבות תדמית, ראיונות טלוויזיה / רדיו, אזכורים ב־Forbes / Globes / Ynet וכו׳ — "
        "קישורים או סריקות. יוצרים סעיף «הזיכו אותנו» עדין בפוטר.",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_field(doc, "קישור 1:")
    add_field(doc, "קישור 2:")
    add_field(doc, "קישור 3:")

    # =================================================================
    # SECTION 5 — TESTIMONIALS & TRUST
    # =================================================================
    add_eyebrow(doc, "חלק 5 · אמינות וביקורות")
    add_heading(doc, "Social proof — מה שעושה רושם ברמה יוקרתית", level=1)
    add_rule(doc)

    add_heading(doc, "5.1  ציטוטי לקוחות", level=2)
    add_p(
        doc,
        "3–5 ציטוטים אמתיים של לקוחות יוצרים אפקט חזק של אמון. אפשר גם קצרים (שורה). "
        "חשוב: לא ציטוטים המצאוניים. אם אין — נשתמש רק בלינק לגוגל.",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    for i in range(1, 4):
        add_field(doc, f"ציטוט #{i} (עד 2 שורות):")
        add_field(doc, f"שם הלקוח / עיר / תחום (רק שם פרטי או ראשי תיבות מספיק):")

    add_heading(doc, "5.2  Case Studies / עסקאות מייצגות (אופציונלי)", level=2)
    add_p(
        doc,
        "2–3 עסקאות שמוצגות עם רמת פירוט גבוהה: רקע הלקוח, האתגר, מה הפתרון של שיבומי היה, "
        "התוצאה. בלי מחירים מדויקים אם לא רוצים. דף נפרד באתר.",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_checkbox_line(doc, "כן, נעבוד על 2–3 case studies בהמשך")
    add_checkbox_line(doc, "לא, עדיף להישאר עם הפורטפוליו הרגיל")

    add_heading(doc, "5.3  לקוחות ש־NDA / עסקאות דיסקרטיות", level=2)
    add_p(
        doc,
        "אם יש נכסים / עסקאות שאסור להזכיר באתר — נא ציין אותם, נוציא מהמלאי:",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_field(doc, "")

    # =================================================================
    # SECTION 6 — ADMIN SYSTEM
    # =================================================================
    add_eyebrow(doc, "חלק 6 · מערכת ניהול (Admin)")
    add_heading(doc, "הגדרות אבטחה ושימוש", level=1)
    add_rule(doc)

    add_heading(doc, "6.1  מי נכנס למערכת?", level=2)
    add_p(
        doc,
        "כל משתמש במערכת צריך שם משתמש וסיסמה משלו. מי המשתמשים המורשים?",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_field(doc, "משתמש 1 (שם מלא / תפקיד):")
    add_field(doc, "משתמש 2 (שם מלא / תפקיד):")
    add_field(doc, "משתמש 3 (שם מלא / תפקיד):")
    add_p(
        doc,
        "(ברירת מחדל: משתמש יחיד של מייקל. נוסיף משתמשים רק לפי בקשה.)",
        size=10, italic=True, color=(120, 120, 130), after_pt=2,
    )

    add_heading(doc, "6.2  שחזור סיסמה", level=2)
    add_p(
        doc,
        "לפי ההנחיה שלך — שחזור סיסמה מתבצע ידנית דרך שיחה ל־050-733-0090. "
        "מי עונה לטלפון הזה עבור בקשות שחזור?",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_field(doc, "שם וגורם מוסמך לאפס סיסמאות:")

    add_heading(doc, "6.3  Two-Factor Authentication (רצוי)", level=2)
    add_p(
        doc,
        "מעבר לסיסמה חזקה, אפשר להוסיף קוד חד־פעמי מאפליקציה (Google Authenticator / Authy). "
        "מחזק משמעותית את האבטחה. הוספה היא עבודה של יום.",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_checkbox_line(doc, "כן, נוסיף 2FA")
    add_checkbox_line(doc, "לא לעכשיו, סיסמה חזקה מספיקה")

    add_heading(doc, "6.4  גיבוי נתונים", level=2)
    add_p(
        doc,
        "באיזו תדירות לגבות את רשימת הנכסים / פניות / סטטיסטיקות?",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_checkbox_line(doc, "גיבוי אוטומטי יומי למייל של מייקל")
    add_checkbox_line(doc, "גיבוי שבועי")
    add_checkbox_line(doc, "ידני לפי בקשה בלבד")

    # =================================================================
    # SECTION 7 — LEGAL
    # =================================================================
    add_eyebrow(doc, "חלק 7 · משפטי ומדיניות")
    add_heading(doc, "דפים חובה על פי חוק", level=1)
    add_rule(doc)

    add_heading(doc, "7.1  הצהרת נגישות", level=2)
    add_p(
        doc,
        "נוסח הצהרת נגישות סטנדרטית ישראלית. נשתמש באותו שירות «נגיש בקליק» שמופיע באתר הקיים, "
        "או בתבנית עצמאית משלנו. נא לאשר:",
        size=11, color=(60, 63, 75), after_pt=4,
    )
    add_checkbox_line(doc, "להמשיך עם «נגיש בקליק» — יש לי חשבון / חוזה פעיל")
    add_checkbox_line(doc, "להשתמש בפתרון חלופי (לסגור איתכם)")

    add_heading(doc, "7.2  מדיניות פרטיות ותקנון", level=2)
    add_checkbox_line(doc, "יש לי מסמכים קיימים — אשלח")
    add_checkbox_line(doc, "כתבו אתם טיוטה ל־GDPR-friendly בעברית")

    add_heading(doc, "7.3  פרטי חברה לפוטר", level=2)
    add_field(doc, "שם חברה רשום מלא:")
    add_field(doc, "מספר עוסק / ח.פ:")
    add_field(doc, "כתובת למשלוח דואר רשמי:")

    # =================================================================
    # APPENDIX — WHAT WE DECIDED (so owner knows)
    # =================================================================
    add_eyebrow(doc, "נספח · החלטות עיצוב שקיבלנו")
    add_heading(doc, "לידיעתך, כדי שלא תופתע", level=1)
    add_rule(doc)

    add_heading(doc, "א · טיפוגרפיה", level=2)
    add_p(
        doc,
        "הפונטים שנבחרו עבור האתר — אחרי בחינה מול שפת המותג של שיבומי ומול הנסיבות של עברית + לטינית "
        "משלבות יחד:",
        size=11, color=(60, 63, 75), after_pt=6,
    )
    add_p(doc, "Fraunces (Latin Display)", size=12, bold=True, after_pt=1)
    add_p(
        doc,
        "serif מודרני של הסטודיו Undercase. תכונה מיוחדת: אופטיקה משתנה (text grade, soft axis) — "
        "הכותרות בגדלים גדולים מקבלות אופי חם ורהוט, הטקסטים הקטנים נשארים קלאסיים. "
        "נדיר מאוד באתרי נדל״ן ישראלים — זה יוצר חתימה.",
        size=10, color=(90, 97, 110), after_pt=6,
    )

    add_p(doc, "Frank Ruhl Libre (Hebrew Display + Body)", size=12, bold=True, after_pt=1)
    add_p(
        doc,
        "פונט עברי קלאסי-יוקרתי של Yanek Iontef, משומש על ידי עיתון «הארץ». "
        "משלב מאוד טוב עם Fraunces הלטיני, ומחזיק את האיכות הסיפורית שהשפה דורשת.",
        size=10, color=(90, 97, 110), after_pt=6,
    )

    add_p(doc, "Assistant (Hebrew UI / Body Secondary)", size=12, bold=True, after_pt=1)
    add_p(
        doc,
        "sans-serif עברי נקי של Ort Type לטקסטים פונקציונליים (טופסים, כפתורים, metadata). "
        "נבחר על פני Heebo כי הוא פחות שכיח ומדויק יותר באותיות קטנות.",
        size=10, color=(90, 97, 110), after_pt=6,
    )

    add_p(doc, "JetBrains Mono (Micro Labels / Numbers)", size=12, bold=True, after_pt=1)
    add_p(
        doc,
        "monospace לצגיוץ מידות (דונמים, מ״ר, שנה, מספרים ארכיוניים) — יוצר תחושה של architectural drawing.",
        size=10, color=(90, 97, 110), after_pt=6,
    )

    add_heading(doc, "ב · פלטת צבעים", level=2)
    add_p(
        doc,
        "Ink 900 (אנתרציט עמוק · #14171C) · Cream 50 (קרם חם · #FBF9F4) · "
        "Brass 600 (זהב מעושן עדין · #9A8551) בלבד. "
        "אין גרדיאנטים, אין צבעי AI. המותג משדר באמצעות אוויר וניגודיות ולא באמצעות צבע.",
        size=11, color=(60, 63, 75), after_pt=6,
    )

    add_heading(doc, "ג · כיוון ויזואלי", level=2)
    add_p(
        doc,
        "Wabi × Pragma — מינימליזם יפני פגוש ברציונליות שוויצרית. "
        "הקאנג׳י 渋 מופיע כסימן־מים עדין בנקודות מפתח, טיפוגרפיה גדולה נחתכת "
        "על ידי תמונות בצורה editorial, כל נכס מקבל מספר ארכיוני ומדידות בשוליים. "
        "רחוק במודע מהסטנדרט של אתרי נדל״ן ישראלים.",
        size=11, color=(60, 63, 75), after_pt=12,
    )

    add_rule(doc)

    add_p(
        doc,
        "סיום המסמך.  לכל שאלה / הערה — ניתן להשיב בכל קטע, או להתקשר ישירות לצוות הפיתוח.",
        size=10, italic=True, color=(120, 120, 130), align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, after_pt=2,
    )

    out = Path(__file__).parent / "Shibumi_Nadlan_Brief_To_Owner.docx"
    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    build()
